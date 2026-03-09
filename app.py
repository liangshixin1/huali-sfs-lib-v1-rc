import os
import pandas as pd
import io
import json
import time
import secrets
import requests
import re
import barcode
from barcode.writer import ImageWriter
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from flask import Flask, render_template, request, redirect, url_for, flash, abort, Response, stream_with_context, jsonify, session, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, logout_user, current_user, login_required
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from sqlalchemy import or_, func, desc, asc
from sqlalchemy.exc import NoResultFound
from datetime import datetime, timedelta, date
from openai import OpenAI
import jieba
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
from dotenv import load_dotenv
from functools import wraps
from urllib.parse import urlencode, urlparse
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

load_dotenv()

# --- 基本配置 ---
basedir = os.path.abspath(os.path.dirname(__file__))
app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'a_default_fallback_secret_key')
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + os.path.join(basedir, 'instance', 'library.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['REMEMBER_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['SESSION_COOKIE_SECURE'] = os.getenv('FLASK_ENV') == 'production'
app.config['MAX_CONTENT_LENGTH'] = int(os.getenv('MAX_CONTENT_LENGTH', 10 * 1024 * 1024))
app.config['ALLOWED_EXCEL_EXTENSIONS'] = {'xlsx'}
app.jinja_env.autoescape = True

# --- AI & Speech API 配置 ---
AI_DETAILS_API_KEY = os.getenv("AI_DETAILS_API_KEY")
details_client = OpenAI(api_key=AI_DETAILS_API_KEY, base_url="https://api.deepseek.com")
AI_SEARCH_API_KEY = os.getenv("AI_SEARCH_API_KEY")
search_client = OpenAI(api_key=AI_SEARCH_API_KEY, base_url="https://api.deepseek.com")
WORD_VECTOR_API_KEY = os.getenv("WORD_VECTOR_API_KEY")
vector_client = OpenAI(api_key=WORD_VECTOR_API_KEY, base_url="https://api.deepseek.com")
BAIDU_APP_ID = os.getenv("BAIDU_SPEECH_APP_ID")
BAIDU_API_KEY = os.getenv("BAIDU_SPEECH_API_KEY")
BAIDU_SECRET_KEY = os.getenv("BAIDU_SPEECH_SECRET_KEY")

# --- RAG & Baidu Token Cache ---
vectorizer, book_vectors, book_corpus_data = None, None, []
baidu_token_cache = {'token': None, 'expires_at': 0}

retry_strategy = Retry(total=3, backoff_factor=1, status_forcelist=[500, 502, 503, 504])
_safe_http_session = requests.Session()
adapter = HTTPAdapter(max_retries=retry_strategy)
_safe_http_session.mount('http://', adapter)
_safe_http_session.mount('https://', adapter)
ALLOWED_EXTERNAL_HOSTS = {'aip.baidubce.com', 'tsn.baidu.com'}


def safe_api_request(url, method='post', **kwargs):
    """Perform outbound HTTP requests with SSRF and timeout protections."""
    parsed = urlparse(url)
    if parsed.scheme not in {'http', 'https'}:
        raise ValueError('Unsupported URL scheme')
    if parsed.hostname not in ALLOWED_EXTERNAL_HOSTS:
        raise ValueError('Disallowed external host')

    kwargs.setdefault('timeout', 10)
    kwargs.setdefault('allow_redirects', False)
    response = _safe_http_session.request(method.upper(), url, **kwargs)
    response.raise_for_status()
    return response


def is_allowed_excel_file(file_storage):
    """Validate uploaded Excel files via extension, header and parsing checks."""
    if not file_storage or not file_storage.filename:
        return False

    filename = secure_filename(file_storage.filename)
    if not filename.lower().endswith('.xlsx'):
        return False

    file_storage.stream.seek(0)
    file_header = file_storage.stream.read(4)
    file_storage.stream.seek(0)
    if not file_header.startswith(b'PK'):
        return False

    try:
        pd.read_excel(file_storage, nrows=0)
        file_storage.stream.seek(0)
    except Exception:
        return False

    return True


def escape_like_specials(value: str) -> str:
    return value.replace('\\', '\\\\').replace('%', '\\%').replace('_', '\\_')


def build_fuzzy_filters(columns, raw_value):
    if not raw_value:
        return []
    sanitized = escape_like_specials(raw_value.strip())
    if not sanitized:
        return []
    pattern = f"%{sanitized}%"
    return [column.ilike(pattern, escape='\\') for column in columns]


def sanitize_excel_cell(value):
    if isinstance(value, str):
        if value.startswith(('=', '+', '-', '@')) or value.startswith('\t') or value.startswith('\r'):
            return "'" + value
    return value


def sanitize_book_payload(data):
    allowed_fields = {'title', 'publisher', 'isbn', 'book_code', 'stock', 'bookshelf_number'}
    sanitized = {}
    for field in allowed_fields:
        if field in data:
            value = data.get(field)
            if isinstance(value, str):
                value = value.strip()
            sanitized[field] = value

    if 'stock' in sanitized:
        try:
            sanitized['stock'] = max(0, int(sanitized['stock']))
        except (TypeError, ValueError):
            raise ValueError('库存量必须为非负整数。')
    if sanitized.get('bookshelf_number') == '':
        sanitized['bookshelf_number'] = None
    if sanitized.get('publisher') == '':
        sanitized['publisher'] = None
    return sanitized


def generate_csrf_token():
    token = session.get('_csrf_token')
    if not token:
        token = secrets.token_urlsafe(32)
        session['_csrf_token'] = token
    return token


app.jinja_env.globals['csrf_token'] = generate_csrf_token


@app.before_request
def enforce_csrf_protection():
    if request.method in {'POST', 'PUT', 'PATCH', 'DELETE'}:
        session_token = session.get('_csrf_token')
        submitted_token = request.headers.get('X-CSRFToken')
        if not request.is_json:
            submitted_token = request.form.get('_csrf_token') or submitted_token
        if not session_token or not submitted_token or not secrets.compare_digest(session_token, submitted_token):
            abort(400)


# --- 数据库和登录管理器初始化 ---
db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'
login_manager.login_message = '请先登录以访问此页面。'
login_manager.login_message_category = 'info'


@app.after_request
def apply_security_headers(response):
    response.headers.setdefault('X-Content-Type-Options', 'nosniff')
    response.headers.setdefault('X-Frame-Options', 'SAMEORIGIN')
    response.headers.setdefault('Referrer-Policy', 'strict-origin-when-cross-origin')
    response.headers.setdefault('Permissions-Policy', 'camera=(self), microphone=(), geolocation=()')

    csp = (
        "default-src 'self'; "
        "script-src 'self' 'unsafe-inline' https://cdn.tailwindcss.com https://cdnjs.cloudflare.com; "
        "style-src 'self' 'unsafe-inline' https://fonts.googleapis.com https://cdnjs.cloudflare.com; "
        "font-src 'self' https://fonts.gstatic.com https://cdnjs.cloudflare.com; "
        "img-src 'self' data:; "
        "connect-src 'self'; "
        "frame-ancestors 'self'"
    )
    response.headers.setdefault('Content-Security-Policy', csp)

    if request.is_secure:
        response.headers.setdefault('Strict-Transport-Security', 'max-age=31536000; includeSubDomains')

    return response


@app.context_processor
def inject_current_year(): return {'current_year': datetime.utcnow().year}

# --- 数据库模型定义 ---
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    full_name = db.Column(db.String(100), nullable=True)
    password_hash = db.Column(db.String(256), nullable=False)
    role = db.Column(db.String(20), nullable=False, default='faculty') 
    records = db.relationship('BorrowRecord', backref='borrower', lazy=True, cascade="all, delete-orphan")
    wants = db.relationship('Want', backref='user', lazy='dynamic', cascade="all, delete-orphan")
    api_usages = db.relationship('ApiUsage', backref='user', lazy='dynamic', cascade="all, delete-orphan")
    def set_password(self, password): self.password_hash = generate_password_hash(password)
    def check_password(self, password): return check_password_hash(self.password_hash, password)
    def has_wanted(self, book): return self.wants.filter_by(book_id=book.id).count() > 0

class Book(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False, index=True)
    publisher = db.Column(db.String(200), nullable=True)
    bookshelf_number = db.Column(db.String(50), nullable=True, index=True)
    isbn = db.Column(db.String(20), nullable=False, index=True)  # 移除unique约束
    book_code = db.Column(db.String(50), unique=True, nullable=False, index=True)
    stock = db.Column(db.Integer, default=1)
    quantity_available = db.Column(db.Integer, default=1)
    records = db.relationship('BorrowRecord', backref='book', lazy=True, cascade="all, delete-orphan")

class BorrowRecord(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    book_id = db.Column(db.Integer, db.ForeignKey('book.id'), nullable=False)
    borrow_date = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)
    due_date = db.Column(db.DateTime, nullable=False, default=lambda: datetime.utcnow() + timedelta(days=30))
    return_date = db.Column(db.DateTime)
    status = db.Column(db.String(20), nullable=False, default='borrowed')

class Want(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    book_id = db.Column(db.Integer, db.ForeignKey('book.id'), nullable=False)
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)

class ApiUsage(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    api_name = db.Column(db.String(50), nullable=False)
    usage_date = db.Column(db.Date, nullable=False, default=date.today)
    count = db.Column(db.Integer, default=0)
    __table_args__ = (db.UniqueConstraint('user_id', 'api_name', 'usage_date', name='_user_api_date_uc'),)


# --- 图书贴纸与上架功能模型 ---
class PendingBook(db.Model):
    """待上架新书表 - 管理员上传的待处理图书"""
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False, index=True)
    publisher = db.Column(db.String(200), nullable=True)
    isbn = db.Column(db.String(20), nullable=False, index=True)
    is_series = db.Column(db.Boolean, default=False)  # 是否为丛书
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    is_active = db.Column(db.Boolean, default=True)  # 是否激活（未被完全入库）


class ShelvingTask(db.Model):
    """上架工作任务表 - 师生的工作记录"""
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    status = db.Column(db.String(20), default='in_progress')  # in_progress, printed, completed
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    items = db.relationship('ShelvingItem', backref='task', lazy='dynamic', cascade="all, delete-orphan")
    user = db.relationship('User', backref='shelving_tasks')


class ShelvingItem(db.Model):
    """上架条目表 - 具体的图书上架记录"""
    id = db.Column(db.Integer, primary_key=True)
    task_id = db.Column(db.Integer, db.ForeignKey('shelving_task.id'), nullable=False)
    pending_book_id = db.Column(db.Integer, db.ForeignKey('pending_book.id'), nullable=False)
    book_code = db.Column(db.String(50), nullable=False, index=True)  # 自动生成的图书编码
    series_count = db.Column(db.Integer, default=1)  # 丛书本数（非丛书为1）
    bookshelf_row = db.Column(db.Integer, nullable=True)  # 第几排
    bookshelf_col = db.Column(db.Integer, nullable=True)  # 第几柜
    is_finalized = db.Column(db.Boolean, default=False)  # 是否已正式入库
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    pending_book = db.relationship('PendingBook', backref='shelving_items')

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

def chinese_tokenizer(text):
    return jieba.lcut(text)


# --- 图书编码生成辅助函数 ---
def get_next_book_code():
    """获取当前年份的下一个可用图书编码（格式：WY年份四位数）"""
    current_year = datetime.utcnow().year
    year_suffix = str(current_year)[2:]  # 取年份后两位
    prefix = f"WY{year_suffix}"

    # 查找数据库中该年份的最大编码
    # 需要同时查找Book表和ShelvingItem表（未入库的）
    max_code_in_books = db.session.query(func.max(Book.book_code)).filter(
        Book.book_code.like(f"{prefix}%")
    ).scalar()

    max_code_in_shelving = db.session.query(func.max(ShelvingItem.book_code)).filter(
        ShelvingItem.book_code.like(f"{prefix}%"),
        ShelvingItem.is_finalized == False
    ).scalar()

    def extract_number(code):
        """从编码中提取基础数字部分"""
        if not code:
            return 0
        # 移除前缀WYxx
        remainder = code[4:]
        # 提取主编号（可能包含-或/）
        match = re.match(r'^(\d+)', remainder)
        if match:
            return int(match.group(1))
        return 0

    max_num_books = extract_number(max_code_in_books)
    max_num_shelving = extract_number(max_code_in_shelving)

    next_num = max(max_num_books, max_num_shelving) + 1
    return f"{prefix}{next_num:04d}"


def get_copy_number(isbn):
    """获取指定ISBN的下一个副本号"""
    # 查找该ISBN在Book表中已存在的数量
    existing_count_in_books = Book.query.filter_by(isbn=isbn).count()

    # 查找该ISBN在ShelvingItem表中已存在的数量（未入库的）
    existing_count_in_shelving = db.session.query(ShelvingItem).join(PendingBook).filter(
        PendingBook.isbn == isbn,
        ShelvingItem.is_finalized == False
    ).count()

    total = existing_count_in_books + existing_count_in_shelving
    return total + 1 if total > 0 else 1


def generate_book_codes_for_item(pending_book, series_count=1):
    """为一个上架条目生成图书编码

    Args:
        pending_book: PendingBook实例
        series_count: 丛书本数（非丛书为1）

    Returns:
        list: 生成的编码列表
    """
    codes = []
    copy_num = get_copy_number(pending_book.isbn)

    if pending_book.is_series and series_count > 1:
        # 丛书情况
        base_code = get_next_book_code()
        for i in range(1, series_count + 1):
            if copy_num > 1:
                codes.append(f"{base_code}-{i}/{copy_num}")
            else:
                codes.append(f"{base_code}-{i}")
    else:
        # 单本书情况
        base_code = get_next_book_code()
        if copy_num > 1:
            codes.append(f"{base_code}/{copy_num}")
        else:
            codes.append(base_code)

    return codes


def generate_barcode_image(code):
    """生成条形码图片并返回BytesIO对象"""
    Code128 = barcode.get_barcode_class('code128')
    rv = io.BytesIO()
    code128 = Code128(code, writer=ImageWriter())
    code128.write(rv, options={
        'module_width': 0.4,
        'module_height': 15,
        'font_size': 10,
        'text_distance': 5,
        'quiet_zone': 2
    })
    rv.seek(0)
    return rv

# --- API 用量限制装饰器 ---
def limit_api_usage(api_name, limit=15):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if current_user.is_anonymous or current_user.role == 'admin':
                return f(*args, **kwargs)
            today = date.today()
            usage = current_user.api_usages.filter_by(api_name=api_name, usage_date=today).first()
            if usage and usage.count >= limit:
                return jsonify({'error': f'您今日的"{api_name}"服务调用次数已达上限。'}), 429
            if not usage:
                usage = ApiUsage(user_id=current_user.id, api_name=api_name, usage_date=today, count=1)
                db.session.add(usage)
            else:
                usage.count += 1
            db.session.commit()
            return f(*args, **kwargs)
        return decorated_function
    return decorator


class BorrowError(Exception):
    def __init__(self, message):
        super().__init__(message)
        self.message = message


def borrow_book_with_lock(book_id, user_id):
    try:
        book = (
            db.session.query(Book)
            .with_for_update()
            .filter_by(id=book_id)
            .one()
        )
        if book.quantity_available <= 0:
            raise BorrowError('该书已无库存可借。')
        book.quantity_available -= 1
        db.session.add(BorrowRecord(user_id=user_id, book_id=book.id))
        book_title = book.title
        db.session.commit()
        return True, f'成功借阅《{book_title}》。', 'success'
    except BorrowError as exc:
        db.session.rollback()
        return False, exc.message, 'warning'
    except NoResultFound:
        db.session.rollback()
        return False, '未找到该书籍。', 'danger'
    except Exception:
        db.session.rollback()
        app.logger.exception('Unexpected error when borrowing book %s for user %s', book_id, user_id)
        return False, '借阅失败，请稍后重试。', 'danger'


# --- 百度语音服务 ---
def get_baidu_token():
    now = time.time()
    if baidu_token_cache['token'] and baidu_token_cache['expires_at'] > now:
        return baidu_token_cache['token']
    url = "https://aip.baidubce.com/oauth/2.0/token"
    params = {"grant_type": "client_credentials", "client_id": BAIDU_API_KEY, "client_secret": BAIDU_SECRET_KEY}
    try:
        response = safe_api_request(url, method='post', params=params)
        data = response.json()
    except (requests.RequestException, ValueError) as exc:
        app.logger.warning('Failed to obtain Baidu token: %s', exc)
        return None
    baidu_token_cache['token'] = data.get("access_token")
    baidu_token_cache['expires_at'] = now + data.get("expires_in", 3600) - 60
    return baidu_token_cache['token']

# --- 搜索建议API ---
@app.route('/search-suggestions')
@login_required
def search_suggestions():
    query = request.args.get('q', '').strip()
    if not query or len(query) < 1:
        return jsonify([])

    # 搜索书名和编码
    filters = build_fuzzy_filters([Book.title, Book.book_code, Book.publisher], query)
    if not filters:
        return jsonify([])

    books = Book.query.filter(or_(*filters)).limit(10).all()
    
    suggestions = []
    seen = set()
    
    for book in books:
        # 书名建议
        if book.title not in seen:
            suggestions.append({
                'text': book.title,
                'type': 'title',
                'icon': 'book'
            })
            seen.add(book.title)
        
        # 编码建议
        if book.book_code not in seen:
            suggestions.append({
                'text': book.book_code,
                'type': 'code',
                'icon': 'code'
            })
            seen.add(book.book_code)
    
    return jsonify(suggestions[:8])

# --- 所有路由 ---
@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated: return redirect(url_for('admin_dashboard' if current_user.role == 'admin' else 'index'))
    if request.method == 'POST':
        login_type, username, password = request.form.get('type'), request.form.get('username'), request.form.get('password')
        user = User.query.filter_by(username=username).first()
        if user and user.check_password(password) and ((login_type == 'admin' and user.role == 'admin') or (login_type == 'faculty' and user.role == 'faculty')):
            session.clear()
            login_user(user)
            return redirect(url_for('admin_dashboard' if user.role == 'admin' else 'index'))
        else: flash('用户名、密码或登录类型错误。', 'danger')
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('您已成功退出登录。', 'success')
    return redirect(url_for('login'))

@app.route('/')
@login_required
def index():
    if current_user.role == 'admin': return redirect(url_for('admin_dashboard'))
    page = request.args.get('page', 1, type=int)
    search_query = request.args.get('search_query', '')
    sort_by = request.args.get('sort_by', 'id')
    sort_order = request.args.get('sort_order', 'desc')

    query = Book.query
    if search_query:
        filters = build_fuzzy_filters(
            [Book.title, Book.book_code, Book.publisher, Book.bookshelf_number],
            search_query
        )
        if filters:
            query = query.filter(or_(*filters))

    # 排序逻辑
    sort_column = getattr(Book, sort_by, Book.id)
    if sort_order == 'desc':
        query = query.order_by(desc(sort_column))
    else:
        query = query.order_by(asc(sort_column))
    
    books_pagination = query.paginate(page=page, per_page=10, error_out=False)
    return render_template('index.html', books_pagination=books_pagination, search_query=search_query, sort_by=sort_by, sort_order=sort_order)

@app.route('/my-records')
@login_required
def my_records():
    if current_user.role != 'faculty': abort(403)
    page = request.args.get('page', 1, type=int)
    borrowed_count = BorrowRecord.query.filter_by(user_id=current_user.id, status='borrowed').count()
    all_records_pagination = BorrowRecord.query.filter_by(user_id=current_user.id).order_by(BorrowRecord.borrow_date.desc()).paginate(page=page, per_page=10, error_out=False)
    return render_template('records.html', records_pagination=all_records_pagination, borrowed_count=borrowed_count)

@app.route('/borrow/<int:book_id>', methods=['POST'])
@login_required
def borrow_book(book_id):
    if current_user.role != 'faculty': abort(403)
    book = Book.query.get_or_404(book_id)
    existing_loan = BorrowRecord.query.filter_by(user_id=current_user.id, book_id=book.id, status='borrowed').first()
    if existing_loan:
        flash(f'您已借阅《{book.title}》，无法重复借阅。', 'warning')
    else:
        success, message, category = borrow_book_with_lock(book.id, current_user.id)
        flash(message, category)
    return redirect(request.referrer or url_for('index'))

@app.route('/borrow-by-code/<string:book_code>', methods=['POST'])
@login_required
def borrow_by_code(book_code):
    if current_user.role != 'faculty': return jsonify({'success': False, 'message': '权限不足'}), 403
    book = Book.query.filter_by(book_code=book_code).first()
    if not book: return jsonify({'success': False, 'message': '未找到该书'}), 404
    existing_loan = BorrowRecord.query.filter_by(user_id=current_user.id, book_id=book.id, status='borrowed').first()
    if existing_loan: return jsonify({'success': False, 'message': f'您已借阅《{book.title}》'})
    success, message, category = borrow_book_with_lock(book.id, current_user.id)
    if success:
        return jsonify({'success': True, 'message': message})
    status_code = 400 if category == 'warning' else 500
    return jsonify({'success': False, 'message': message}), status_code


@app.route('/return-by-code/<string:book_code>', methods=['POST'])
@login_required
def return_by_code(book_code):
    if current_user.role != 'faculty':
        return jsonify({'success': False, 'message': '权限不足'}), 403

    record = (
        BorrowRecord.query
        .join(Book)
        .filter(
            BorrowRecord.user_id == current_user.id,
            BorrowRecord.status == 'borrowed',
            Book.book_code == book_code
        )
        .first()
    )

    if not record:
        return jsonify({'success': False, 'message': '未找到与该编码匹配的在借记录'}), 404

    record.status = 'returned'
    record.return_date = datetime.utcnow()
    record.book.quantity_available += 1
    db.session.commit()

    return jsonify({'success': True, 'message': f'成功归还《{record.book.title}》'})

@app.route('/want/<int:book_id>', methods=['POST'])
@login_required
def want_book(book_id):
    if current_user.role != 'faculty': abort(403)
    book = Book.query.get_or_404(book_id)
    if not current_user.has_wanted(book):
        db.session.add(Want(user_id=current_user.id, book_id=book.id))
        db.session.commit()
        flash(f'已将《{book.title}》加入您的想读列表。', 'success')
    else: flash(f'《{book.title}》已在您的想读列表中。', 'info')
    return redirect(request.referrer or url_for('index'))

@app.route('/return/<int:record_id>', methods=['POST'])
@login_required
def return_book(record_id):
    if current_user.role != 'faculty': abort(403)
    record = BorrowRecord.query.get_or_404(record_id)
    if record.user_id != current_user.id: abort(403)
    if record.status == 'borrowed':
        record.status, record.return_date = 'returned', datetime.utcnow()
        record.book.quantity_available += 1
        db.session.commit()
        flash(f'成功归还《{record.book.title}》。', 'success')
    else: flash('该记录状态异常，无法归还。', 'warning')
    return redirect(url_for('my_records'))

@app.route('/return-all', methods=['POST'])
@login_required
def return_all_books():
    if current_user.role != 'faculty': abort(403)
    records_to_return = BorrowRecord.query.filter_by(user_id=current_user.id, status='borrowed').all()
    if not records_to_return: flash('没有需要归还的书籍。', 'info')
    else:
        for record in records_to_return:
            record.status, record.return_date = 'returned', datetime.utcnow()
            record.book.quantity_available += 1
        db.session.commit()
        flash(f'成功归还 {len(records_to_return)} 本书籍。', 'success')
    return redirect(url_for('my_records'))

@app.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    if current_user.role != 'faculty': abort(403)
    if request.method == 'POST':
        old_password, new_password, confirm_password = request.form.get('old_password'), request.form.get('new_password'), request.form.get('confirm_password')
        if not current_user.check_password(old_password): flash('当前密码不正确。', 'danger')
        elif new_password != confirm_password: flash('两次输入的新密码不一致。', 'danger')
        elif len(new_password) < 4: flash('新密码长度不能少于4位。', 'danger')
        else:
            current_user.set_password(new_password)
            db.session.commit()
            logout_user()
            flash('密码修改成功，请重新登录。', 'success')
            return redirect(url_for('login'))
    return render_template('profile.html')

@app.route('/ask-ai', methods=['POST'])
@login_required
@limit_api_usage('AI问答')
def ask_ai():
    data = request.get_json()
    chat_history = data.get('history', [])
    if not chat_history: return Response(json.dumps({'error': 'History is empty'}), status=400, mimetype='application/json')
    def generate():
        try:
            stream = details_client.chat.completions.create(model="deepseek-chat", messages=chat_history, stream=True)
            for chunk in stream:
                if content := chunk.choices[0].delta.content: yield content
        except Exception as e:
            print(f"AI Details API Error: {e}")
            yield "抱歉，图书详情AI服务暂时无法连接。"
    return Response(stream_with_context(generate()), mimetype='text/event-stream')

@app.route('/ai-search', methods=['POST'])
@login_required
@limit_api_usage('AI搜书')
def ai_search():
    data = request.get_json()
    user_query = data.get('query')
    expanded_query = data.get('expanded_query')
    if not user_query: return Response(json.dumps({'error': 'Query is empty'}), status=400, mimetype='application/json')
    if vectorizer is None or book_vectors is None: return Response(json.dumps({'error': 'AI Search engine not initialized.'}), status=503)
    final_query = expanded_query if expanded_query else user_query
    query_vector = vectorizer.transform([final_query])
    similarities = cosine_similarity(query_vector, book_vectors).flatten()
    top_n_indices = np.argsort(similarities)[-5:][::-1]
    retrieved_context = [f"- 书名:《{book['title']}》, 作者:{book['publisher']}, ISBN:{book['isbn']}, 状态:{'可用' if book['quantity_available'] > 0 else '已借出'}, 图书编码:{book['book_code']}" for index in top_n_indices if similarities[index] > 0.01 and (book := book_corpus_data[index])]
    context_str = "\n".join(retrieved_context) if retrieved_context else "未在书库中找到直接相关的书籍。"
    system_prompt = f"""你是一个专业的图书查询助手。请严格根据下面提供的"相关书籍资料"来回答用户的问题。
相关书籍资料:
---
{context_str}
---
请根据以上资料，回答用户的问题："{user_query}"。如果资料中有合适的书籍，请推荐给用户并简要说明理由。如果资料显示未找到相关书籍，请如实告知用户。请保持回答简洁。"""
    messages = [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_query}]
    def generate():
        try:
            stream = search_client.chat.completions.create(model="deepseek-chat", messages=messages, stream=True)
            for chunk in stream:
                if content := chunk.choices[0].delta.content: yield content
        except Exception as e:
            print(f"AI Search API Error: {e}")
            yield "抱歉，AI搜书服务在生成回答时遇到问题。"
    return Response(stream_with_context(generate()), mimetype='text/event-stream')

@app.route('/ai-expand-query', methods=['POST'])
@login_required
@limit_api_usage('AI语义扩展')
def ai_expand_query():
    data = request.get_json()
    query = data.get('query', '').strip()
    if not query: return jsonify({'success': False, 'message': 'Query is empty'}), 400
    system_prompt = """你是一个专为图书馆检索系统服务的、高效的语义解析引擎。你的唯一任务是将用户的自然语言查询，严格地转换为一个结构化的JSON对象，以便后续的关键词匹配。严格遵守以下处理规则：
1. 核心实体识别: 识别查询中的核心概念，如人物、事件、主题、作品等。
2. 同义词/近义词扩展: 为核心实体提供2-3个最相关的同义或近义词。
3. 时间范围解析: 如果查询中包含时期描述（如"维多利亚时期"、"明末清初"），将其解析为一个大致的年份范围（如 "1837-1901"）。
4. 生成上位/下位概念: 为核心主题生成1个更宽泛的父概念（如"英国文学"）和1-2个更具体的子概念（如"工业小说"）。
5. 补充学术关键词: 基于查询意图，补充3-5个最可能相关的学术关键词（比如：相关作家、学者、作品、理论、术语和学术关键词）。
输出格式: 必须严格返回一个JSON对象。"""
    try:
        response = vector_client.chat.completions.create(model="deepseek-chat", messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": query}], response_format={"type": "json_object"})
        expanded_data = json.loads(response.choices[0].message.content)
        print("\n--- Word_Vector API Response ---")
        print(json.dumps(expanded_data, indent=2, ensure_ascii=False))
        print("----------------------------\n")
        all_terms = set()
        def extract_terms(data):
            if isinstance(data, dict):
                for key, value in data.items():
                    extract_terms(value)
            elif isinstance(data, list):
                for item in data:
                    extract_terms(item)
            elif isinstance(data, str):
                all_terms.add(data)
        extract_terms(expanded_data)
        return jsonify({'success': True, 'expanded_query': " ".join(all_terms)})
    except Exception as e:
        print(f"AI Query Expansion Error: {e}")
        return jsonify({'success': False, 'message': 'AI查询扩展失败'}), 500

@app.route('/ai-recommend', methods=['POST'])
@login_required
def ai_recommend():
    data = request.get_json()
    query = data.get('query', '').strip()
    if not query: return jsonify({'success': False, 'message': 'Query is empty'}), 400
    if vectorizer is None or book_vectors is None: return jsonify({'success': False, 'message': 'RAG engine not ready'}), 503
    query_vector = vectorizer.transform([query])
    similarities = cosine_similarity(query_vector, book_vectors).flatten()
    top_n_indices = np.argsort(similarities)[-5:][::-1]
    recommendations = [{'title': book_corpus_data[i]['title'], 'code': book_corpus_data[i]['book_code']} for i in top_n_indices if similarities[i] > 0.01 and book_corpus_data[i]['quantity_available'] > 0]
    return jsonify({'success': True, 'recommendations': recommendations})

@app.route('/text-to-speech', methods=['POST'])
@login_required
@limit_api_usage('语音合成')
def text_to_speech():
    text = request.json.get('text', '')
    if not text: return jsonify({'error': 'No text provided'}), 400
    token = get_baidu_token()
    if not token: return jsonify({'error': 'Could not get access token'}), 500
    payload = {
        'tex': text,  # 直接使用文本
        'tok': token, 
        'cuid': current_user.username, 
        'ctp': 1, 
        'lan': 'zh', 
        'spd': 5, 
        'pit': 5, 
        'vol': 9, 
        'per': 4226, 
        'aue': 3
    }
    
    try:
        response = safe_api_request("https://tsn.baidu.com/text2audio", method='post', data=payload)
    except (requests.RequestException, ValueError) as exc:
        app.logger.warning('Text-to-speech request failed: %s', exc)
        return jsonify({'error': '语音合成服务暂时不可用'}), 502

    if response.headers.get("Content-Type") == "audio/mp3":
        return Response(response.content, mimetype='audio/mp3')
    try:
        return jsonify(response.json()), response.status_code
    except ValueError:
        return jsonify({'error': '语音合成服务返回了无法解析的结果'}), 502

# --- 图书贴纸与上架路由 ---
@app.route('/shelving')
@login_required
def shelving_page():
    """师生端上架工作页面"""
    # 获取用户当前未完成的任务
    current_task = ShelvingTask.query.filter_by(
        user_id=current_user.id,
        status='in_progress'
    ).first()

    # 如果有已打印但未完成的任务，也显示
    if not current_task:
        current_task = ShelvingTask.query.filter_by(
            user_id=current_user.id,
            status='printed'
        ).first()

    return render_template('shelving.html', current_task=current_task)


@app.route('/shelving/search')
@login_required
def shelving_search():
    """搜索待上架图书"""
    query = request.args.get('q', '').strip()
    if not query or len(query) < 1:
        return jsonify([])

    # 搜索待上架的书籍（按书名模糊搜索或ISBN后四位）
    results = []
    pending_books = PendingBook.query.filter(
        PendingBook.is_active == True
    )

    # 书名模糊搜索
    title_filters = build_fuzzy_filters([PendingBook.title], query)
    if title_filters:
        title_results = pending_books.filter(or_(*title_filters)).limit(10).all()
        results.extend(title_results)

    # ISBN搜索（支持任意四位数）
    if len(query) >= 4 and query.isdigit():
        isbn_results = pending_books.filter(
            PendingBook.isbn.contains(query)
        ).limit(10).all()
        for book in isbn_results:
            if book not in results:
                results.append(book)

    return jsonify([{
        'id': book.id,
        'title': book.title,
        'publisher': book.publisher or '',
        'isbn': book.isbn,
        'is_series': book.is_series
    } for book in results[:10]])


@app.route('/shelving/start', methods=['POST'])
@login_required
def shelving_start():
    """开始新的上架工作"""
    # 检查是否有未完成的任务
    existing_task = ShelvingTask.query.filter_by(
        user_id=current_user.id
    ).filter(ShelvingTask.status.in_(['in_progress', 'printed'])).first()

    if existing_task:
        return jsonify({
            'success': True,
            'task_id': existing_task.id,
            'message': '已恢复您之前的工作'
        })

    # 创建新任务
    new_task = ShelvingTask(user_id=current_user.id)
    db.session.add(new_task)
    db.session.commit()

    return jsonify({
        'success': True,
        'task_id': new_task.id,
        'message': '已开始新的上架工作'
    })


@app.route('/shelving/add-book', methods=['POST'])
@login_required
def shelving_add_book():
    """添加图书到待上架列表"""
    data = request.get_json()
    pending_book_id = data.get('pending_book_id')
    series_count = data.get('series_count', 1)

    if not pending_book_id:
        return jsonify({'success': False, 'message': '请选择要添加的图书'}), 400

    pending_book = PendingBook.query.get(pending_book_id)
    if not pending_book or not pending_book.is_active:
        return jsonify({'success': False, 'message': '图书不存在或已下架'}), 404

    # 获取或创建任务
    task = ShelvingTask.query.filter_by(
        user_id=current_user.id
    ).filter(ShelvingTask.status.in_(['in_progress', 'printed'])).first()

    if not task:
        task = ShelvingTask(user_id=current_user.id)
        db.session.add(task)
        db.session.commit()

    # 使用事务锁确保编码唯一性
    try:
        with db.session.begin_nested():
            # 生成图书编码
            base_code = get_next_book_code()
            copy_num = get_copy_number(pending_book.isbn)

            # 处理丛书
            if pending_book.is_series and series_count > 1:
                codes = []
                for i in range(1, series_count + 1):
                    if copy_num > 1:
                        codes.append(f"{base_code}-{i}/{copy_num}")
                    else:
                        codes.append(f"{base_code}-{i}")
                book_code = ';'.join(codes)
            else:
                if copy_num > 1:
                    book_code = f"{base_code}/{copy_num}"
                else:
                    book_code = base_code

            # 创建上架条目
            item = ShelvingItem(
                task_id=task.id,
                pending_book_id=pending_book_id,
                book_code=book_code,
                series_count=series_count
            )
            db.session.add(item)

        db.session.commit()

        return jsonify({
            'success': True,
            'item': {
                'id': item.id,
                'title': pending_book.title,
                'publisher': pending_book.publisher or '',
                'isbn': pending_book.isbn,
                'book_code': book_code,
                'series_count': series_count,
                'is_series': pending_book.is_series
            }
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': f'添加失败: {str(e)}'}), 500


@app.route('/shelving/remove-item/<int:item_id>', methods=['POST'])
@login_required
def shelving_remove_item(item_id):
    """移除上架条目"""
    item = ShelvingItem.query.get_or_404(item_id)

    # 验证权限
    if item.task.user_id != current_user.id:
        return jsonify({'success': False, 'message': '无权操作'}), 403

    if item.is_finalized:
        return jsonify({'success': False, 'message': '已入库的图书无法移除'}), 400

    db.session.delete(item)
    db.session.commit()

    return jsonify({'success': True, 'message': '已移除'})


@app.route('/shelving/get-items')
@login_required
def shelving_get_items():
    """获取当前任务的所有条目"""
    task = ShelvingTask.query.filter_by(
        user_id=current_user.id
    ).filter(ShelvingTask.status.in_(['in_progress', 'printed'])).first()

    if not task:
        return jsonify({'items': []})

    items = []
    for item in task.items.filter_by(is_finalized=False).all():
        items.append({
            'id': item.id,
            'title': item.pending_book.title,
            'publisher': item.pending_book.publisher or '',
            'isbn': item.pending_book.isbn,
            'book_code': item.book_code,
            'series_count': item.series_count,
            'is_series': item.pending_book.is_series,
            'bookshelf_row': item.bookshelf_row,
            'bookshelf_col': item.bookshelf_col
        })

    return jsonify({
        'task_id': task.id,
        'status': task.status,
        'items': items
    })


@app.route('/shelving/print-stickers', methods=['POST'])
@login_required
def shelving_print_stickers():
    """生成贴纸文档（5x2格式）"""
    task = ShelvingTask.query.filter_by(
        user_id=current_user.id
    ).filter(ShelvingTask.status.in_(['in_progress', 'printed'])).first()

    if not task:
        return jsonify({'success': False, 'message': '没有进行中的任务'}), 400

    items = task.items.filter_by(is_finalized=False).all()
    if not items:
        return jsonify({'success': False, 'message': '没有待打印的图书'}), 400

    # 收集所有需要打印的贴纸数据
    stickers = []
    for item in items:
        codes = item.book_code.split(';')
        for code in codes:
            stickers.append({
                'title': item.pending_book.title,
                'publisher': item.pending_book.publisher or '',
                'code': code.strip()
            })

    # 创建Word文档
    document = Document()

    # 设置页面边距
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    # 计算需要多少个表格（每个表格5行2列）
    stickers_per_page = 10
    num_pages = (len(stickers) + stickers_per_page - 1) // stickers_per_page

    for page_num in range(num_pages):
        start_idx = page_num * stickers_per_page
        end_idx = min(start_idx + stickers_per_page, len(stickers))
        page_stickers = stickers[start_idx:end_idx]

        # 创建5x2表格
        table = document.add_table(rows=5, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 设置表格宽度
        for row in table.rows:
            row.height = Cm(5)
            for cell in row.cells:
                cell.width = Cm(9)

        # 填充贴纸内容
        sticker_idx = 0
        for row_idx in range(5):
            for col_idx in range(2):
                if sticker_idx < len(page_stickers):
                    sticker = page_stickers[sticker_idx]
                    cell = table.cell(row_idx, col_idx)

                    # 清空单元格
                    cell.text = ''

                    # 添加书名
                    p_title = cell.add_paragraph()
                    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run_title = p_title.add_run(sticker['title'][:40] + ('...' if len(sticker['title']) > 40 else ''))
                    run_title.bold = True
                    run_title.font.size = Pt(10)

                    # 添加出版社
                    p_pub = cell.add_paragraph()
                    p_pub.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run_pub = p_pub.add_run(sticker['publisher'][:30] if sticker['publisher'] else '')
                    run_pub.font.size = Pt(8)

                    # 生成条形码并添加
                    try:
                        barcode_img = generate_barcode_image(sticker['code'])
                        p_barcode = cell.add_paragraph()
                        p_barcode.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_barcode = p_barcode.add_run()
                        run_barcode.add_picture(barcode_img, width=Cm(6))
                    except Exception as e:
                        # 如果条形码生成失败，只显示编码文字
                        p_code = cell.add_paragraph()
                        p_code.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_code = p_code.add_run(sticker['code'])
                        run_code.bold = True
                        run_code.font.size = Pt(14)

                    # 添加编码文字
                    p_code_text = cell.add_paragraph()
                    p_code_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_code_text = p_code_text.add_run(sticker['code'])
                    run_code_text.bold = True
                    run_code_text.font.size = Pt(12)

                sticker_idx += 1

        # 如果不是最后一页，添加分页符
        if page_num < num_pages - 1:
            document.add_page_break()

    # 更新任务状态
    task.status = 'printed'
    db.session.commit()

    # 保存文档到内存
    docx_buffer = io.BytesIO()
    document.save(docx_buffer)
    docx_buffer.seek(0)

    # 返回文件
    filename = f"book_stickers_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return send_file(
        docx_buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )


@app.route('/shelving/update-bookshelf', methods=['POST'])
@login_required
def shelving_update_bookshelf():
    """更新书柜号"""
    data = request.get_json()
    item_id = data.get('item_id')
    row = data.get('row')
    col = data.get('col')

    item = ShelvingItem.query.get_or_404(item_id)

    if item.task.user_id != current_user.id:
        return jsonify({'success': False, 'message': '无权操作'}), 403

    if item.is_finalized:
        return jsonify({'success': False, 'message': '已入库的图书无法修改'}), 400

    item.bookshelf_row = row if row else None
    item.bookshelf_col = col if col else None
    db.session.commit()

    return jsonify({'success': True, 'message': '书柜号已更新'})


@app.route('/shelving/finalize', methods=['POST'])
@login_required
def shelving_finalize():
    """正式入库"""
    task = ShelvingTask.query.filter_by(
        user_id=current_user.id
    ).filter(ShelvingTask.status.in_(['in_progress', 'printed'])).first()

    if not task:
        return jsonify({'success': False, 'message': '没有进行中的任务'}), 400

    items = task.items.filter_by(is_finalized=False).all()
    if not items:
        return jsonify({'success': False, 'message': '没有待入库的图书'}), 400

    finalized_count = 0
    for item in items:
        # 生成书柜号
        if item.bookshelf_row and item.bookshelf_col:
            bookshelf_number = f"{item.bookshelf_row}排{item.bookshelf_col}柜"
        else:
            bookshelf_number = "待定"

        # 处理每个编码
        codes = item.book_code.split(';')
        for code in codes:
            code = code.strip()
            # 创建Book记录
            new_book = Book(
                title=item.pending_book.title,
                publisher=item.pending_book.publisher,
                isbn=item.pending_book.isbn,
                book_code=code,
                stock=1,
                quantity_available=1,
                bookshelf_number=bookshelf_number
            )
            db.session.add(new_book)
            finalized_count += 1

        # 标记为已入库
        item.is_finalized = True

    # 更新任务状态
    task.status = 'completed'
    db.session.commit()

    # 更新AI目录
    update_ai_catalog_on_startup()

    return jsonify({
        'success': True,
        'message': f'成功入库 {finalized_count} 本书籍',
        'count': finalized_count
    })


@app.route('/shelving/save', methods=['POST'])
@login_required
def shelving_save():
    """保存工作（不入库）"""
    task = ShelvingTask.query.filter_by(
        user_id=current_user.id
    ).filter(ShelvingTask.status.in_(['in_progress', 'printed'])).first()

    if not task:
        return jsonify({'success': False, 'message': '没有进行中的任务'}), 400

    # 任务状态保持不变，只是确认保存
    db.session.commit()

    return jsonify({'success': True, 'message': '工作已保存，下次可继续'})


@app.route('/shelving/clear-task', methods=['POST'])
@login_required
def shelving_clear_task():
    """清除当前任务（重新开始）"""
    task = ShelvingTask.query.filter_by(
        user_id=current_user.id
    ).filter(ShelvingTask.status.in_(['in_progress', 'printed'])).first()

    if task:
        # 删除未入库的条目
        task.items.filter_by(is_finalized=False).delete()

        # 如果任务没有已入库的条目，删除任务
        if task.items.count() == 0:
            db.session.delete(task)

    db.session.commit()
    return jsonify({'success': True, 'message': '已清除当前任务'})


# --- 管理员上架管理路由 ---
@app.route('/admin/pending-books')
@login_required
def admin_pending_books():
    """管理员查看待上架图书列表"""
    if current_user.role != 'admin':
        abort(403)

    pending_books = PendingBook.query.filter_by(is_active=True).order_by(PendingBook.created_at.desc()).all()
    return render_template('admin_pending_books.html', pending_books=pending_books)


@app.route('/admin/import-pending-books', methods=['POST'])
@login_required
def admin_import_pending_books():
    """管理员导入待上架新书列表"""
    if current_user.role != 'admin':
        abort(403)

    file = request.files.get('pending_book_file')
    if not file or file.filename == '':
        flash('未选择任何文件。', 'danger')
        return redirect(url_for('admin_dashboard'))

    if not is_allowed_excel_file(file):
        flash('请上传受支持的 .xlsx 格式文件。', 'danger')
        return redirect(url_for('admin_dashboard'))

    try:
        df = pd.read_excel(file)
        required_columns = {'书名', 'ISBN'}
        if not required_columns.issubset(df.columns):
            flash(f'Excel文件必须包含以下列: {", ".join(required_columns)}', 'danger')
            return redirect(url_for('admin_dashboard'))

        added_count = 0
        for _, row in df.iterrows():
            title = str(row['书名']).strip()
            isbn = str(row['ISBN']).strip()

            if not title or not isbn or title == 'nan' or isbn == 'nan':
                continue

            publisher = str(row.get('出版社', '')).strip() if '出版社' in df.columns else ''
            if publisher == 'nan':
                publisher = ''

            # 检查是否为丛书
            is_series = False
            if '是否为丛书' in df.columns:
                series_val = str(row.get('是否为丛书', '')).strip().lower()
                is_series = series_val in ['是', '1', 'true', 'yes', 'y']

            # 创建待上架记录
            pending_book = PendingBook(
                title=title,
                publisher=publisher,
                isbn=isbn,
                is_series=is_series
            )
            db.session.add(pending_book)
            added_count += 1

        db.session.commit()
        flash(f'成功导入 {added_count} 本待上架新书。', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'导入失败: {e}', 'danger')

    return redirect(url_for('admin_dashboard'))


@app.route('/admin/pending-book/delete/<int:book_id>', methods=['POST'])
@login_required
def admin_delete_pending_book(book_id):
    """删除待上架图书"""
    if current_user.role != 'admin':
        abort(403)

    book = PendingBook.query.get_or_404(book_id)
    book.is_active = False
    db.session.commit()
    flash(f'已下架《{book.title}》', 'success')
    return redirect(url_for('admin_dashboard'))


@app.route('/admin/pending-bookshelf')
@login_required
def admin_pending_bookshelf():
    """管理员查看待定书柜号的图书"""
    if current_user.role != 'admin':
        abort(403)

    pending_books = Book.query.filter_by(bookshelf_number='待定').order_by(Book.id.desc()).all()
    return jsonify([{
        'id': book.id,
        'title': book.title,
        'book_code': book.book_code,
        'publisher': book.publisher or ''
    } for book in pending_books])


@app.route('/admin/batch-update-bookshelf', methods=['POST'])
@login_required
def admin_batch_update_bookshelf():
    """批量更新书柜号"""
    if current_user.role != 'admin':
        return jsonify({'success': False, 'message': '无权操作'}), 403

    data = request.get_json()
    updates = data.get('updates', [])

    updated_count = 0
    for update in updates:
        book_id = update.get('id')
        row = update.get('row')
        col = update.get('col')

        book = Book.query.get(book_id)
        if book:
            if row and col:
                book.bookshelf_number = f"{row}排{col}柜"
            else:
                book.bookshelf_number = "待定"
            updated_count += 1

    db.session.commit()

    return jsonify({
        'success': True,
        'message': f'成功更新 {updated_count} 本书的书柜号',
        'count': updated_count
    })


# --- 管理员路由 ---
@app.route('/admin/dashboard')
@login_required
def admin_dashboard():
    if current_user.role != 'admin': abort(403)
    books = Book.query.order_by(Book.id.desc()).all()
    faculty_users = User.query.filter_by(role='faculty').order_by(User.id).all()
    wants_query = db.session.query(Book.title, func.count(Want.id).label('want_count')).join(Want).group_by(Book.title).order_by(func.count(Want.id).desc()).all()
    filters = {'name': request.args.get('name', ''), 'start_date': request.args.get('start_date', ''), 'end_date': request.args.get('end_date', ''), 'status': request.args.get('status', 'all')}
    query = BorrowRecord.query.join(User).join(Book)
    if filters['name']:
        name_filters = build_fuzzy_filters([User.full_name], filters['name'])
        if name_filters:
            query = query.filter(name_filters[0])
    if filters['start_date']: query = query.filter(BorrowRecord.borrow_date >= datetime.strptime(filters['start_date'], '%Y-%m-%d'))
    if filters['end_date']: query = query.filter(BorrowRecord.borrow_date < (datetime.strptime(filters['end_date'], '%Y-%m-%d') + timedelta(days=1)))
    if filters['status'] != 'all': query = query.filter(BorrowRecord.status == filters['status'])
    page = request.args.get('page', 1, type=int)
    records_pagination = query.order_by(BorrowRecord.borrow_date.desc()).paginate(page=page, per_page=15, error_out=False)
    # 获取待上架图书列表
    pending_books = PendingBook.query.filter_by(is_active=True).order_by(PendingBook.created_at.desc()).all()
    return render_template('admin_dashboard.html', books=books, faculty_users=faculty_users, wants_list=wants_query, records_pagination=records_pagination, filters=filters, pending_books=pending_books)

@app.route('/admin/update-ai-catalog', methods=['POST'])
@login_required
def update_ai_catalog():
    if current_user.role != 'admin': abort(403)
    try:
        global vectorizer, book_vectors, book_corpus_data
        all_books = Book.query.all()
        if not all_books:
            flash('书库为空，无法更新AI目录。', 'warning')
            return redirect(url_for('admin_dashboard'))
        corpus = [f"{b.title} {b.publisher} {b.bookshelf_number}" for b in all_books]
        book_corpus_data = [{'title': b.title, 'publisher': b.publisher, 'isbn': b.isbn, 'book_code': b.book_code, 'quantity_available': b.quantity_available, 'bookshelf_number': b.bookshelf_number} for b in all_books]
        vectorizer = TfidfVectorizer(tokenizer=chinese_tokenizer)
        book_vectors = vectorizer.fit_transform(corpus)
        flash(f'AI检索引擎更新成功！共索引 {len(all_books)} 本书。', 'success')
    except Exception as e:
        flash(f'更新AI检索引擎时发生错误: {e}', 'danger')
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/import-books', methods=['POST'])
@login_required
def import_books():
    if current_user.role != 'admin': abort(403)
    file = request.files.get('book_file')
    upload_mode = request.form.get('upload_mode', 'append')

    if not file or file.filename == '':
        flash('未选择任何文件。', 'danger')
        return redirect(url_for('admin_dashboard'))

    if not is_allowed_excel_file(file):
        flash('请上传受支持的 .xlsx 格式文件。', 'danger')
        return redirect(url_for('admin_dashboard'))

    try:
        if upload_mode == 'overwrite':
            db.session.query(BorrowRecord).delete()
            db.session.query(Want).delete()
            db.session.query(Book).delete()
            db.session.commit()
            flash('已清空现有书库，正在导入新数据...', 'info')

        df = pd.read_excel(file)
        required_columns = {'书名', '出版社', 'ISBN', '编码', '库存量', '书柜号'}
        if not required_columns.issubset(df.columns):
            flash(f'Excel文件必须包含以下列: {", ".join(required_columns)}', 'danger')
            return redirect(url_for('admin_dashboard'))

        added_count = 0
        skipped_count = 0

        for _, row in df.iterrows():
            book_code, isbn = str(row['编码']).strip(), str(row['ISBN']).strip()

            try:
                stock = int(row['库存量'])
            except (ValueError, TypeError):
                skipped_count += 1
                continue

            existing_book = Book.query.filter_by(book_code=book_code).first()

            if not existing_book:
                db.session.add(Book(
                    title=row['书名'],
                    publisher=row['出版社'],
                    isbn=isbn,
                    book_code=book_code,
                    stock=stock,
                    quantity_available=stock,
                    bookshelf_number=row.get('书柜号')
                ))
                added_count += 1
            else:
                skipped_count += 1

        db.session.commit()
        flash(f'数据导入完成！新增 {added_count} 本书，跳过 {skipped_count} 本重复编码的书。', 'success')
        update_ai_catalog_on_startup()
        flash('AI检索引擎已同步更新。', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'导入失败: {e}', 'danger')
    return redirect(url_for('admin_dashboard'))


@app.route('/admin/import-users', methods=['POST'])
@login_required
def admin_import_users():
    if current_user.role != 'admin':
        abort(403)

    file = request.files.get('user_file')
    if not file or file.filename == '':
        flash('未选择任何用户数据文件。', 'danger')
        return redirect(url_for('admin_dashboard'))

    if not is_allowed_excel_file(file):
        flash('请上传受支持的 .xlsx 格式用户文件。', 'danger')
        return redirect(url_for('admin_dashboard'))

    try:
        df = pd.read_excel(file)
    except Exception as exc:
        flash(f'无法读取上传的用户文件: {exc}', 'danger')
        return redirect(url_for('admin_dashboard'))

    required_sets = [
        {'工号', '教职工姓名', '初始密码'},
        {'工号', '教职工姓名', '身份证后四位'}
    ]

    if not any(columns.issubset(df.columns) for columns in required_sets):
        flash('Excel文件必须包含「工号」「教职工姓名」以及「初始密码」或「身份证后四位」列。', 'danger')
        return redirect(url_for('admin_dashboard'))

    password_column = '初始密码' if '初始密码' in df.columns else '身份证后四位'

    added, skipped = 0, 0

    for _, row in df.iterrows():
        username = str(row.get('工号', '')).strip()
        full_name = str(row.get('教职工姓名', '')).strip() or username
        raw_password = row.get(password_column)

        password = None
        if pd.isna(raw_password):
            password = None
        elif isinstance(raw_password, float):
            if raw_password.is_integer():
                password = str(int(raw_password))
            else:
                password = str(raw_password)
        elif isinstance(raw_password, (int, np.integer)):
            password = str(raw_password)
        elif isinstance(raw_password, str):
            password = raw_password.strip()

        if password:
            password = password.strip()
            if password.lower() == 'nan':
                password = None

        if not username or not password:
            skipped += 1
            continue

        if User.query.filter_by(username=username).first():
            skipped += 1
            continue

        new_user = User(
            username=username,
            full_name=full_name,
            role='faculty',
            password_hash=generate_password_hash(password)
        )
        db.session.add(new_user)
        added += 1

    if added:
        db.session.commit()
        flash(f'成功导入 {added} 位新用户，跳过 {skipped} 条记录。', 'success')
    else:
        db.session.rollback()
        flash(f'没有导入新的用户，跳过 {skipped} 条记录。', 'warning')

    return redirect(url_for('admin_dashboard'))

@app.route('/admin/return', methods=['POST'])
@login_required
def admin_return_book():
    if current_user.role != 'admin': abort(403)
    book_code = request.form.get('book_code')
    record = BorrowRecord.query.join(Book).filter(Book.book_code == book_code, BorrowRecord.status == 'borrowed').first()
    if not record: flash('未找到该图书编码对应的已借出记录。', 'danger')
    else:
        record.status, record.return_date = 'returned', datetime.utcnow()
        record.book.quantity_available += 1
        db.session.commit()
        flash(f'《{record.book.title}》已成功归还。', 'success')
    return redirect(url_for('admin_dashboard'))
    
@app.route('/admin/book/add', methods=['POST'])
@login_required
def admin_add_book():
    if current_user.role != 'admin': abort(403)
    try:
        book_data = sanitize_book_payload(request.form)
    except ValueError as exc:
        flash(str(exc), 'danger')
        return redirect(url_for('admin_dashboard'))

    required_fields = {'title', 'isbn', 'book_code', 'stock'}
    if not required_fields.issubset(book_data.keys()):
        flash('请完整填写书名、ISBN、图书编码和库存量。', 'danger')
        return redirect(url_for('admin_dashboard'))

    if not all(book_data.get(field) for field in {'title', 'isbn', 'book_code'}):
        flash('书名、ISBN 和图书编码不能为空。', 'danger')
        return redirect(url_for('admin_dashboard'))

    if Book.query.filter_by(book_code=book_data['book_code']).first():
        flash('图书编码已存在。', 'danger')
    else:
        new_book = Book(
            title=book_data['title'],
            publisher=book_data.get('publisher'),
            isbn=book_data['isbn'],
            book_code=book_data['book_code'],
            stock=book_data['stock'],
            quantity_available=book_data['stock'],
            bookshelf_number=book_data.get('bookshelf_number')
        )
        db.session.add(new_book)
        db.session.commit()
        flash(f'书籍《{new_book.title}》添加成功。', 'success')
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/book/update/<int:book_id>', methods=['POST'])
@login_required
def admin_update_book(book_id):
    if current_user.role != 'admin': abort(403)
    book = Book.query.get_or_404(book_id)
    try:
        incoming = sanitize_book_payload(request.form)
    except ValueError as exc:
        flash(str(exc), 'danger')
        return redirect(url_for('admin_dashboard'))

    if 'book_code' in incoming and incoming['book_code'] != book.book_code and Book.query.filter_by(book_code=incoming['book_code']).first():
        flash('图书编码已存在。', 'danger')
        return redirect(url_for('admin_dashboard'))

    stock_change = 0
    if 'stock' in incoming:
        stock_change = incoming['stock'] - book.stock

    for required in ('title', 'isbn', 'book_code'):
        if required in incoming and not incoming[required]:
            flash('书名、ISBN 和图书编码不能为空。', 'danger')
            return redirect(url_for('admin_dashboard'))

    for field in ('title', 'publisher', 'isbn', 'book_code', 'bookshelf_number'):
        if field in incoming:
            setattr(book, field, incoming[field])

    if 'stock' in incoming:
        book.stock = incoming['stock']
        book.quantity_available = max(0, min(book.quantity_available + stock_change, book.stock))

    db.session.commit()
    flash(f'书籍《{book.title}》信息更新成功。', 'success')
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/book/delete/<int:book_id>', methods=['POST'])
@login_required
def admin_delete_book(book_id):
    if current_user.role != 'admin': abort(403)
    book = Book.query.get_or_404(book_id)
    flash(f'书籍《{book.title}》已删除。', 'success')
    db.session.delete(book)
    db.session.commit()
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/user/add', methods=['POST'])
@login_required
def admin_add_user():
    if current_user.role != 'admin': abort(403)
    username, full_name, password = request.form.get('username'), request.form.get('full_name'), request.form.get('password')
    if not all([username, full_name, password]): flash('所有字段均为必填项。', 'danger')
    elif User.query.filter_by(username=username).first(): flash('该工号已被注册。', 'danger')
    else:
        db.session.add(User(username=username, full_name=full_name, role='faculty', password_hash=generate_password_hash(password)))
        db.session.commit()
        flash(f'用户 {full_name} ({username}) 添加成功。', 'success')
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/user/update/<int:user_id>', methods=['POST'])
@login_required
def admin_update_user(user_id):
    if current_user.role != 'admin': abort(403)
    user = User.query.get_or_404(user_id)
    user.full_name = request.form.get('full_name')
    if new_password := request.form.get('password'):
        user.set_password(new_password)
        flash(f'用户 {user.full_name} 的姓名和密码已更新。', 'success')
    else:
        flash(f'用户 {user.full_name} 的姓名已更新。', 'success')
    db.session.commit()
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/user/delete/<int:user_id>', methods=['POST'])
@login_required
def admin_delete_user(user_id):
    if current_user.role != 'admin': abort(403)
    user = User.query.get_or_404(user_id)
    if user.role == 'admin':
        flash('不能删除管理员账户。', 'danger')
        return redirect(url_for('admin_dashboard'))
    flash(f'用户 {user.full_name} ({user.username}) 已被成功删除。', 'success')
    db.session.delete(user)
    db.session.commit()
    return redirect(url_for('admin_dashboard'))
    
@app.route('/admin/export-records')
@login_required
def admin_export_records():
    if current_user.role != 'admin': abort(403)
    filters = {'name': request.args.get('name', ''), 'start_date': request.args.get('start_date', ''), 'end_date': request.args.get('end_date', ''), 'status': request.args.get('status', 'all')}
    query = BorrowRecord.query.join(User).join(Book)
    if filters['name']:
        name_filters = build_fuzzy_filters([User.full_name], filters['name'])
        if name_filters:
            query = query.filter(name_filters[0])
    if filters['start_date']: query = query.filter(BorrowRecord.borrow_date >= datetime.strptime(filters['start_date'], '%Y-%m-%d'))
    if filters['end_date']: query = query.filter(BorrowRecord.borrow_date < (datetime.strptime(filters['end_date'], '%Y-%m-%d') + timedelta(days=1)))
    if filters['status'] != 'all': query = query.filter(BorrowRecord.status == filters['status'])
    records = query.order_by(BorrowRecord.borrow_date.desc()).all()
    data = [{'书名': r.book.title, '图书编码': r.book.book_code, '借阅人': r.borrower.full_name, '工号': r.borrower.username, '借阅日期': r.borrow_date.strftime('%Y-%m-%d %H:%M'), '应还日期': r.due_date.strftime('%Y-%m-%d'), '归还日期': r.return_date.strftime('%Y-%m-%d %H:%M') if r.return_date else '未归还', '状态': '已归还' if r.status == 'returned' else '借阅中'} for r in records]
    df = pd.DataFrame(data)
    if not df.empty:
        df = df.applymap(sanitize_excel_cell)
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='借阅记录')
    output.seek(0)
    filename = f"borrow_records_{datetime.now().strftime('%Y%m%d')}.xlsx"
    return Response(output, mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": f"attachment;filename={filename}"})

# --- 应用初始化函数 ---
def init_database():
    with app.app_context():
        db.create_all()
        admin_pass = os.getenv("ADMIN_PASSWORD", "Admin")
        if not User.query.filter_by(username='Admin').first():
            db.session.add(User(username='Admin', full_name='系统管理员', role='admin', password_hash=generate_password_hash(admin_pass)))
            print("Admin user created.")
        
        try:
            user_file = os.path.join(basedir, '账密.xlsx')
            if os.path.exists(user_file):
                df_users = pd.read_excel(user_file)
                for _, row in df_users.iterrows():
                    username = str(row['工号'])
                    if not User.query.filter_by(username=username).first():
                        db.session.add(User(username=username, full_name=row.get('教职工姓名', username), role='faculty', password_hash=generate_password_hash(str(row['身份证后四位']))))
                print(f"Imported users from 账密.xlsx")
        except FileNotFoundError:
            print("Warning: '账密.xlsx' not found. Creating fallback user if defined.")
            fallback_user_data = os.getenv("FALLBACK_FACULTY_USER")
            if fallback_user_data:
                try:
                    username, full_name, password = fallback_user_data.split(',')
                    if not User.query.filter_by(username=username).first():
                        db.session.add(User(username=username, full_name=full_name, role='faculty', password_hash=generate_password_hash(password)))
                        print(f"Created fallback faculty user: {username}")
                except Exception as e:
                    print(f"Error parsing FALLBACK_FACULTY_USER: {e}")

        except Exception as e:
            print(f"Error importing users: {e}")
        
        db.session.commit()
        print("Initializing AI Catalog for the first time...")
        with app.test_request_context():
            update_ai_catalog_on_startup()

def update_ai_catalog_on_startup():
    try:
        global vectorizer, book_vectors, book_corpus_data
        all_books = Book.query.all()
        if not all_books: 
            print("No books in DB, AI engine not initialized.")
            return
        corpus = [f"{b.title} {b.publisher} {b.bookshelf_number}" for b in all_books]
        book_corpus_data = [{'title': b.title, 'publisher': b.publisher, 'isbn': b.isbn, 'book_code': b.book_code, 'quantity_available': b.quantity_available, 'bookshelf_number': b.bookshelf_number} for b in all_books]
        vectorizer = TfidfVectorizer(tokenizer=chinese_tokenizer)
        book_vectors = vectorizer.fit_transform(corpus)
        print(f'AI Search Engine initialized with {len(all_books)} books.')
    except Exception as e:
        print(f"Failed to initialize AI Search Engine: {e}")

# --- 主程序入口 ---
if __name__ == '__main__':
    init_database()
    if not all([AI_DETAILS_API_KEY, AI_SEARCH_API_KEY, WORD_VECTOR_API_KEY, BAIDU_API_KEY, BAIDU_SECRET_KEY]):
        print("\n\033[91m错误: 缺少API Keys！\033[0m")
        print("请确保您已在项目根目录创建了 .env 文件，并包含所有必需的API Keys。\n")
    else:
        app.run(debug=True, threaded=True)